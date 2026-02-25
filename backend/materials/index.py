"""API для работы с учебными материалами: загрузка документов через presigned URL"""

import json
import os
import hashlib
import boto3
from datetime import datetime
import psycopg2
from psycopg2.extras import RealDictCursor
import jwt
from openai import OpenAI
import io
from PyPDF2 import PdfReader
from docx import Document
from rate_limiter import check_rate_limit, get_client_ip
from security_validator import sanitize_filename, check_ownership

MAX_FILE_SIZE = 50 * 1024 * 1024
CHUNK_SIZE = 3500
CHUNK_OVERLAP = 500  # Overlap для сохранения контекста между чанками


def get_db_connection():
    dsn = os.environ['DATABASE_URL']
    schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
    return psycopg2.connect(dsn, options=f'-c search_path={schema}')


def verify_token(token: str) -> dict:
    secret = os.environ['JWT_SECRET']
    try:
        return jwt.decode(token, secret, algorithms=['HS256'])
    except:
        return None


def check_subscription_access(conn, user_id: int) -> dict:
    schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    cursor.execute(f'''
        SELECT subscription_type, subscription_expires_at, trial_ends_at, is_trial_used 
        FROM {schema}.users 
        WHERE id = %s
    ''', (user_id,))
    user = cursor.fetchone()
    cursor.close()
    
    if not user:
        return {'has_access': False, 'reason': 'user_not_found'}
    
    now = datetime.now()
    
    # Проверяем премиум подписку
    if user.get('subscription_type') == 'premium':
        if user.get('subscription_expires_at') and user['subscription_expires_at'].replace(tzinfo=None) > now:
            return {'has_access': True, 'is_premium': True, 'is_trial': False}
    
    # Проверяем триал период
    trial_ends = user.get('trial_ends_at')
    is_trial_used = user.get('is_trial_used')
    
    if trial_ends and not is_trial_used:
        trial_ends_naive = trial_ends.replace(tzinfo=None) if trial_ends.tzinfo else trial_ends
        if trial_ends_naive > now:
            return {'has_access': True, 'is_premium': False, 'is_trial': True, 'trial_ends_at': trial_ends}
    
    return {'has_access': False, 'reason': 'no_subscription'}


def get_s3_client():
    return boto3.client('s3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY'])


def generate_presigned_upload_url(filename: str, file_type: str, user_id: int) -> dict:
    s3 = get_s3_client()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    key = f"materials/{user_id}_{timestamp}_{filename}"
    
    try:
        presigned_url = s3.generate_presigned_url('put_object',
            Params={'Bucket': 'files', 'Key': key, 'ContentType': file_type},
            ExpiresIn=3600)
        
        cdn_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{key}"
        return {'upload_url': presigned_url, 'file_key': key, 'cdn_url': cdn_url}
    except Exception as e:
        print(f"[MATERIALS] Ошибка presigned URL: {e}")
        return None


def download_file_from_s3(file_key: str) -> bytes:
    """Скачивает файл из S3 для обработки на бэкенде"""
    s3 = get_s3_client()
    try:
        print(f"[MATERIALS] Скачиваю из S3: Bucket=files, Key={file_key}")
        response = s3.get_object(Bucket='files', Key=file_key)
        data = response['Body'].read()
        print(f"[MATERIALS] Скачано {len(data)} байт")
        return data
    except Exception as e:
        print(f"[MATERIALS] Ошибка скачивания из S3: {e}")
        import traceback
        traceback.print_exc()
        return None


def extract_text_from_pdf(file_data: bytes) -> str:
    try:
        pdf_reader = PdfReader(io.BytesIO(file_data))
        return '\n\n'.join([page.extract_text() for page in pdf_reader.pages])
    except Exception as e:
        print(f"[MATERIALS] PDF ошибка: {e}")
        return ""


def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        
        # Извлекаем текст из параграфов
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
            if table_text:
                text_parts.append('\n[ТАБЛИЦА]\n' + '\n'.join(table_text) + '\n[/ТАБЛИЦА]\n')
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f"[MATERIALS] DOCX ошибка: {e}")
        return ""


def extract_text_from_txt(file_data: bytes) -> str:
    for encoding in ['utf-8', 'windows-1251', 'cp1251', 'latin-1']:
        try:
            return file_data.decode(encoding)
        except:
            continue
    return file_data.decode('utf-8', errors='ignore')


def extract_text_from_file(file_data: bytes, file_type: str) -> str:
    if 'pdf' in file_type.lower():
        return extract_text_from_pdf(file_data)
    elif 'word' in file_type.lower() or 'document' in file_type.lower():
        return extract_text_from_docx(file_data)
    elif 'text' in file_type.lower() or 'plain' in file_type.lower():
        return extract_text_from_txt(file_data)
    return ""


def split_text_into_chunks(text: str) -> list:
    """Разбивает текст на чанки с overlap для сохранения контекста"""
    if not text:
        return []
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= CHUNK_SIZE:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # Добавляем overlap: берём последние CHUNK_OVERLAP символов
                overlap_text = current_chunk[-CHUNK_OVERLAP:] if len(current_chunk) > CHUNK_OVERLAP else current_chunk
                current_chunk = overlap_text + "\n\n" + para + "\n\n"
            else:
                current_chunk = para + "\n\n"
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks


def analyze_document_with_deepseek(full_text: str, filename: str) -> dict:
    openrouter_key = os.environ.get('OPENROUTER_API_KEY', '')
    
    if not openrouter_key or not full_text or len(full_text) < 10:
        return {'summary': 'Документ загружен', 'subject': 'Общее', 'title': filename[:50], 'tasks': []}
    
    print(f"[MATERIALS] OpenRouter/Llama анализ начат, длина текста={len(full_text)}")
    
    try:
        client = OpenAI(api_key=openrouter_key, base_url="https://api.aitunnel.ru/v1/", timeout=22.0)
        # Берём больше текста для анализа (начало + середина + конец)
        text_start = full_text[:2500]
        text_middle = full_text[len(full_text)//2:len(full_text)//2+2500] if len(full_text) > 5000 else ""
        text_end = full_text[-1000:] if len(full_text) > 3500 else ""
        text_preview = f"{text_start}\n\n[...середина документа...]\n{text_middle}\n\n[...конец документа...]\n{text_end}"
        
        prompt = f"""Ты помощник студента. Проанализируй учебный документ "{filename}".

Фрагменты текста из разных частей документа:
{text_preview}

Верни JSON:
{{"summary": "Подробное резюме документа (5-7 предложений, основные темы и концепции)", "subject": "Предмет", "title": "Название (макс 50 символов)", "tasks": [{{"title": "Задача", "deadline": "YYYY-MM-DD или null"}}]}}"""
        
        response = client.chat.completions.create(
            model="deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            response_format={"type": "json_object"})
        
        content = response.choices[0].message.content
        if '```json' in content:
            content = content.split('```json')[1].split('```')[0].strip()
        elif '```' in content:
            content = content.split('```')[1].split('```')[0].strip()
        
        result = json.loads(content)
        print(f"[MATERIALS] OpenRouter/Llama анализ завершен: title={result.get('title')}, subject={result.get('subject')}")
        return result
    except Exception as e:
        print(f"[MATERIALS] Artemox ошибка: {e}")
        return {'summary': 'Документ загружен (анализ недоступен)', 'subject': 'Общее', 'title': filename[:50], 'tasks': []}


def generate_share_code(material_id: int, user_id: int) -> str:
    """Генерирует детерминированный код для шаринга материала"""
    return hashlib.md5(f"{material_id}:{user_id}".encode()).hexdigest()[:8].upper()


def handle_get_shared(code: str, headers: dict) -> dict:
    """Публичный эндпоинт: получить расшаренный материал по коду"""
    if not code or len(code) != 8:
        return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Неверный код'}, ensure_ascii=False)}

    conn = get_db_connection()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT id, user_id, title, subject, summary, recognized_text FROM materials ORDER BY id")
            for row in cur:
                expected = generate_share_code(row['id'], row['user_id'])
                if expected == code.upper():
                    text = row['recognized_text'] or ''
                    return {
                        'statusCode': 200,
                        'headers': headers,
                        'body': json.dumps({
                            'title': row['title'],
                            'subject': row['subject'],
                            'summary': row['summary'],
                            'recognized_text': text[:5000]
                        }, ensure_ascii=False, default=str)
                    }

        return {'statusCode': 404, 'headers': headers, 'body': json.dumps({'error': 'Материал не найден'}, ensure_ascii=False)}
    finally:
        conn.close()


def handler(event: dict, context) -> dict:
    method = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {'statusCode': 200, 'headers': {'Access-Control-Allow-Origin': '*', 'Access-Control-Allow-Methods': 'GET, POST, DELETE, OPTIONS', 'Access-Control-Allow-Headers': 'Content-Type, Authorization, X-Authorization'}, 'body': ''}
    
    headers = {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'}
    
    # Публичный эндпоинт: просмотр расшаренного материала (без авторизации)
    params = event.get('queryStringParameters') or {}
    if method == 'GET' and params.get('action') == 'shared':
        return handle_get_shared(params.get('code', ''), headers)
    
    auth_header = event.get('headers', {}).get('X-Authorization', '')
    token = auth_header.replace('Bearer ', '')
    
    if not token:
        return {'statusCode': 401, 'headers': headers, 'body': json.dumps({'error': 'Требуется авторизация'})}
    
    payload = verify_token(token)
    if not payload:
        return {'statusCode': 401, 'headers': headers, 'body': json.dumps({'error': 'Недействительный токен'})}
    
    user_id = payload['user_id']
    
    # POST - загрузка файлов
    if method == 'POST':
        body = json.loads(event.get('body', '{}'))
        action = body.get('action')
        
        # Прямая загрузка файла через base64 (обход CORS проблем)
        if action == 'upload_direct':
            try:
                conn = get_db_connection()
                access = check_subscription_access(conn, user_id)
                
                # Проверяем лимит для Free пользователей (3 материала/месяц)
                if not access['has_access']:
                    message = '⏰ Подписка истекла' if access.get('reason') == 'subscription_expired' else '🔒 Требуется подписка'
                    conn.close()
                    return {'statusCode': 403, 'headers': headers, 'body': json.dumps({'error': 'subscription_required', 'message': message})}
                
                # Для Free проверяем месячный лимит (2 материала)
                if not access.get('is_premium') and not access.get('is_trial'):
                    schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
                    with conn.cursor(cursor_factory=RealDictCursor) as cur:
                        cur.execute(f'''
                            SELECT materials_quota_used, materials_quota_reset_at 
                            FROM {schema}.users 
                            WHERE id = %s
                        ''', (user_id,))
                        quota_info = cur.fetchone()
                        
                        # Проверяем, не истек ли месячный лимит
                        quota_used = quota_info.get('materials_quota_used', 0)
                        if quota_used >= 2:
                            conn.close()
                            return {'statusCode': 403, 'headers': headers, 'body': json.dumps({'error': 'quota_exceeded', 'message': '📊 Лимит загрузок исчерпан (2/2). Перейдите на Premium для безлимитных материалов'})}
                
                conn.close()
                
                filename = body.get('filename')
                file_type = body.get('fileType')
                file_data_base64 = body.get('fileData')
                
                if not filename or not file_type or not file_data_base64:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Не указаны filename, fileType или fileData'})}
                
                import base64
                file_data = base64.b64decode(file_data_base64)
                file_size = len(file_data)
                
                if file_size > MAX_FILE_SIZE:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': f'Макс размер: {MAX_FILE_SIZE // 1024 // 1024} МБ'})}
                
                print(f"[MATERIALS] Загрузка {filename}, размер={file_size} байт")
                
                # Загружаем в S3
                s3 = get_s3_client()
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                import random
                file_key = f"materials/{user_id}_{timestamp}_{random.randint(10000000, 99999999)}_{filename}"
                
                s3.put_object(Bucket='files', Key=file_key, Body=file_data, ContentType=file_type)
                cdn_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{file_key}"
                
                print(f"[MATERIALS] Файл загружен в S3: {file_key}")
                
                # Обрабатываем сразу
                print(f"[MATERIALS] Извлекаю текст, тип файла: {file_type}")
                full_text = extract_text_from_file(file_data, file_type)
                
                if not full_text or len(full_text.strip()) < 10:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Файл пуст или не содержит распознаваемого текста'})}
                
                print(f"[MATERIALS] Извлечено {len(full_text)} символов текста")
                
                chunks = split_text_into_chunks(full_text)
                print(f"[MATERIALS] Разбито на {len(chunks)} чанков")
                
                analysis = analyze_document_with_deepseek(full_text, filename)
                print(f"[MATERIALS] DeepSeek результат: {analysis}")
                
                title = (analysis.get('title') or filename)[:200]
                subject = (analysis.get('subject') or 'Общее')[:100]
                summary = (analysis.get('summary') or 'Документ загружен')[:2000]
                file_type_short = file_type[:50]
                
                print(f"[MATERIALS] Данные: title={title[:50]}..., subject={subject}, len(summary)={len(summary)}, file_type={file_type_short}")
                
                conn = get_db_connection()
                print(f"[MATERIALS] БД подключение OK")
                try:
                    with conn.cursor(cursor_factory=RealDictCursor) as cur:
                        print(f"[MATERIALS] Начинаю INSERT materials для user_id={user_id}...")
                        # Для больших документов храним только чанки, для маленьких - весь текст
                        text_preview = full_text[:2000] if len(chunks) > 1 else full_text[:10000]
                        cur.execute("""
                            INSERT INTO materials (user_id, title, subject, file_url, recognized_text, summary, file_type, file_size, total_chunks)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                            RETURNING id, title, subject, file_url, summary, file_type, file_size, total_chunks, created_at
                        """, (user_id, title, subject, cdn_url, text_preview, summary, file_type_short, file_size, len(chunks)))
                        print(f"[MATERIALS] INSERT materials OK")
                        
                        material = cur.fetchone()
                        material_id = material['id']
                        print(f"[MATERIALS] Получен material_id={material_id}")
                        
                        for idx, chunk in enumerate(chunks):
                            cur.execute("INSERT INTO document_chunks (material_id, chunk_index, chunk_text) VALUES (%s, %s, %s)", (material_id, idx, chunk))
                        print(f"[MATERIALS] Вставлено {len(chunks)} чанков")
                        
                        # Увеличиваем счетчик использованных материалов для Free
                        schema = os.environ.get('MAIN_DB_SCHEMA', 'public')
                        cur.execute(f'''
                            UPDATE {schema}.users 
                            SET materials_quota_used = materials_quota_used + 1
                            WHERE id = %s AND subscription_type = 'free'
                        ''', (user_id,))
                        
                        conn.commit()
                        print(f"[MATERIALS] COMMIT OK, материал ID={material_id} создан, квота обновлена")
                        
                        return {'statusCode': 200, 'headers': headers, 'body': json.dumps({'material': dict(material), 'chunks_created': len(chunks)}, default=str)}
                except Exception as db_error:
                    print(f"[MATERIALS] ⚠️ Ошибка БД: {type(db_error).__name__}: {db_error}")
                    import traceback
                    traceback.print_exc()
                    raise
                finally:
                    conn.close()
            except Exception as e:
                print(f"[MATERIALS] ❌ Ошибка upload_direct: {type(e).__name__}: {e}")
                import traceback
                traceback.print_exc()
                return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': str(e)})}
        
        # Шаг 1: получить presigned URL (старый метод, оставляем для совместимости)
        elif action == 'get_upload_url':
            try:
                conn = get_db_connection()
                access = check_subscription_access(conn, user_id)
                conn.close()
                
                if not access['has_access']:
                    message = '⏰ Подписка истекла' if access.get('reason') == 'subscription_expired' else '🔒 Требуется подписка'
                    return {'statusCode': 403, 'headers': headers, 'body': json.dumps({'error': 'subscription_required', 'message': message})}
                
                filename = body.get('filename')
                file_type = body.get('fileType')
                file_size = body.get('fileSize', 0)
                
                if not filename or not file_type:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Не указаны filename и fileType'})}
                
                if file_size > MAX_FILE_SIZE:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': f'Макс размер: {MAX_FILE_SIZE // 1024 // 1024} МБ'})}
                
                presigned_data = generate_presigned_upload_url(filename, file_type, user_id)
                if not presigned_data:
                    return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': 'Ошибка генерации URL'})}
                
                return {'statusCode': 200, 'headers': headers, 'body': json.dumps(presigned_data)}
            except Exception as e:
                print(f"[MATERIALS] Ошибка: {e}")
                import traceback
                traceback.print_exc()
                return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': str(e)})}
        
        # Шаг 2: обработать загруженный файл
        elif action == 'process_file':
            try:
                file_key = body.get('fileKey')
                cdn_url = body.get('cdnUrl')
                filename = body.get('filename')
                file_type = body.get('fileType')
                file_size = body.get('fileSize')
                
                if not file_key or not cdn_url:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Не указаны fileKey/cdnUrl'})}
                
                print(f"[MATERIALS] Обработка файла: {filename}, key={file_key}")
                file_data = download_file_from_s3(file_key)
                
                if not file_data:
                    return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': 'Не удалось скачать файл из хранилища. Попробуйте еще раз.'})}
                
                print(f"[MATERIALS] Извлекаю текст, тип файла: {file_type}")
                full_text = extract_text_from_file(file_data, file_type)
                
                if not full_text or len(full_text.strip()) < 10:
                    return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Файл пуст или не содержит распознаваемого текста'})}
                
                print(f"[MATERIALS] Извлечено {len(full_text)} символов текста")
                
                chunks = split_text_into_chunks(full_text)
                analysis = analyze_document_with_deepseek(full_text, filename)
                
                title = (analysis.get('title') or filename)[:200]
                subject = (analysis.get('subject') or 'Общее')[:100]
                summary = (analysis.get('summary') or 'Документ загружен')[:2000]
                file_type_short = file_type[:50]
                
                conn = get_db_connection()
                try:
                    with conn.cursor(cursor_factory=RealDictCursor) as cur:
                        # Для больших документов храним только чанки, для маленьких - весь текст
                        text_preview = full_text[:2000] if len(chunks) > 1 else full_text[:10000]
                        cur.execute("""
                            INSERT INTO materials (user_id, title, subject, file_url, recognized_text, summary, file_type, file_size, total_chunks)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                            RETURNING id, title, subject, file_url, summary, file_type, file_size, total_chunks, created_at
                        """, (user_id, title, subject, cdn_url, text_preview, summary, file_type_short, file_size, len(chunks)))
                        
                        material = cur.fetchone()
                        material_id = material['id']
                        
                        for idx, chunk in enumerate(chunks):
                            cur.execute("INSERT INTO document_chunks (material_id, chunk_index, chunk_text) VALUES (%s, %s, %s)", (material_id, idx, chunk))
                        
                        conn.commit()
                        print(f"[MATERIALS] Создан: ID={material_id}")
                        return {'statusCode': 201, 'headers': headers, 'body': json.dumps({'material': dict(material), 'tasks': analysis.get('tasks', []), 'chunks_count': len(chunks)}, default=str)}
                finally:
                    conn.close()
            except Exception as e:
                print(f"[MATERIALS] Ошибка: {e}")
                import traceback
                traceback.print_exc()
                return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': str(e)})}
        
        elif action == 'share':
            material_id = body.get('material_id')
            if not material_id:
                return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Не указан material_id'})}
            
            conn = get_db_connection()
            try:
                with conn.cursor() as cur:
                    cur.execute("SELECT id FROM materials WHERE id = %s AND user_id = %s", (material_id, user_id))
                    if not cur.fetchone():
                        return {'statusCode': 404, 'headers': headers, 'body': json.dumps({'error': 'Материал не найден'})}
                
                code = generate_share_code(int(material_id), user_id)
                return {'statusCode': 200, 'headers': headers, 'body': json.dumps({'code': code, 'material_id': material_id}, ensure_ascii=False)}
            finally:
                conn.close()
        
        return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Неизвестное действие'})}
    
    # GET - список материалов
    elif method == 'GET':
        conn = get_db_connection()
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT id, title, subject, file_url, recognized_text, summary, file_type, file_size, total_chunks, created_at FROM materials WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
                materials = cur.fetchall()
                return {'statusCode': 200, 'headers': headers, 'body': json.dumps({'materials': [dict(m) for m in materials]}, default=str)}
        finally:
            conn.close()
    
    # DELETE - удалить материал
    elif method == 'DELETE':
        material_id = event.get('queryStringParameters', {}).get('id')
        if not material_id:
            return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'ID не указан'})}
        
        conn = get_db_connection()
        try:
            with conn.cursor() as cur:
                cur.execute("DELETE FROM document_chunks WHERE material_id IN (SELECT id FROM materials WHERE id = %s AND user_id = %s)", (material_id, user_id))
                cur.execute("DELETE FROM materials WHERE id = %s AND user_id = %s", (material_id, user_id))
                conn.commit()
                
                if cur.rowcount == 0:
                    return {'statusCode': 404, 'headers': headers, 'body': json.dumps({'error': 'Не найден'})}
                
                return {'statusCode': 200, 'headers': headers, 'body': json.dumps({'message': 'Удалён'})}
        finally:
            conn.close()
    
    return {'statusCode': 405, 'headers': headers, 'body': json.dumps({'error': 'Метод не поддерживается'})}